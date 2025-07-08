import streamlit as st
import pandas as pd
from datetime import datetime
import io
import docx
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
import numpy as np
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.cluster import KMeans
from sklearn.metrics.pairwise import cosine_similarity
import re
import json
import base64
from openai import AzureOpenAI
from dotenv import load_dotenv
import os

# Page configuration
st.set_page_config(
    page_title="AI SEO Content Automation Platform",
    page_icon="🚀",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Custom CSS
st.markdown("""
<style>
    .main-header {
        font-size: 2.5rem;
        font-weight: bold;
        color: #1f77b4;
        text-align: center;
        margin-bottom: 2rem;
    }
    .module-header {
        font-size: 1.5rem;
        font-weight: bold;
        color: #2e8b57;
        border-bottom: 2px solid #2e8b57;
        padding-bottom: 0.5rem;
        margin-bottom: 1rem;
    }
    .warning-box {
        background-color: #fff3cd;
        border: 1px solid #ffeaa7;
        border-radius: 0.25rem;
        padding: 1rem;
        margin: 1rem 0;
    }
    .success-box {
        background-color: #d4edda;
        border: 1px solid #c3e6cb;
        border-radius: 0.25rem;
        padding: 1rem;
        margin: 1rem 0;
    }
    .metric-card {
        background-color: #f8f9fa;
        border-radius: 0.5rem;
        padding: 1rem;
        text-align: center;
        border-left: 4px solid #1f77b4;
    }
</style>
""", unsafe_allow_html=True)

# Initialize session state
def init_session_state():
    if 'generated_content' not in st.session_state:
        st.session_state.generated_content = []
    if 'discovered_keywords' not in st.session_state:
        st.session_state.discovered_keywords = []
    if 'keyword_clusters' not in st.session_state:
        st.session_state.keyword_clusters = []
    if 'metadata_results' not in st.session_state:
        st.session_state.metadata_results = []
    if 'api_provider' not in st.session_state:
        st.session_state.api_provider = 'OpenAI'

# Prompt Templates
class PromptTemplates:
    @staticmethod
    def blog_post_template(primary_keyword, secondary_keywords, word_count, tone, industry, context=""):
        return f"""Role: Expert SEO content writer
Task: Create a comprehensive blog post

Requirements:
- Primary keyword: {primary_keyword}
- Secondary keywords: {secondary_keywords}
- Word count: {word_count} words
- Tone: {tone}
- Industry: {industry}

Structure:
1. SEO-optimized title with primary keyword
2. Hook introduction (150 words max)
3. 5-7 main sections with H2 headers
4. Include secondary keywords naturally (2-3% density)
5. Conclusion with CTA

Context: {context}

Format: Return structured markdown with proper H1, H2, H3 tags.
Include a brief meta description at the start."""

    @staticmethod
    def product_page_template(product_name, category, features, keyword, competitor_info=""):
        return f"""Role: E-commerce copywriter specialist
Task: Create product page content

Product Details:
- Name: {product_name}
- Category: {category}
- Key features: {features}
- Target keyword: {keyword}
- Competitor insights: {competitor_info}

Sections needed:
- SEO title & meta description
- Product headline
- Feature bullets (3-5)
- Detailed description (200-300 words)
- FAQ section (3-5 questions)

Focus: Conversion optimization + SEO compliance"""

    @staticmethod
    def metadata_template(content, primary_keyword, search_intent, industry):
        return f"""Role: SEO metadata specialist
Content/Topic: {content[:500]}...

Generate:
1. Title tag (50-60 chars): Include {primary_keyword}
2. Meta description (150-160 chars): Compelling + keyword-rich
3. H1 tag: Primary keyword placement
4. H2 suggestions (3-5): Semantic keywords
5. Image alt text templates

Rules:
- No keyword stuffing
- Include power words for CTR
- Match search intent: {search_intent}
- Industry: {industry}"""

    @staticmethod
    def keyword_discovery_template(seed_keyword, industry, audience, include_local=False, location=""):
        local_instruction = f"Include location-based keywords for {location}" if include_local else ""
        
        return f"""Role: SEO keyword research expert
Seed keyword: {seed_keyword}
Industry: {industry}
Target audience: {audience}

Generate 50 long-tail keywords covering:
1. Question-based (who, what, when, where, why, how)
2. Comparison keywords (vs, compared to, best)
3. Commercial intent (buy, price, cost, cheap)
4. Informational intent (guide, tips, how to)
{local_instruction}

Format each as:
Keyword | Search Intent | Difficulty (1-10) | Content Type

Filter: Only keywords with 2+ words"""

# Utility Functions
def validate_inputs(required_fields, session_state):
    """Validate required inputs and return warnings"""
    warnings = []
    for field, friendly_name in required_fields.items():
        if not session_state.get(field):
            warnings.append(f"💡 {friendly_name} is recommended for better results")
    return warnings

def apply_fallbacks(inputs):
    """Apply smart fallbacks for missing inputs"""
    fallbacks = {
        'industry': 'Technology',
        'tone': 'Professional',
        'word_count': 1500,
        'search_intent': 'Informational'
    }
    
    for key, default in fallbacks.items():
        if not inputs.get(key):
            inputs[key] = default
    
    return inputs

def keyword_clustering(keywords, n_clusters=5):
    """Cluster keywords using TF-IDF and K-means"""
    if len(keywords) < n_clusters:
        n_clusters = len(keywords)
    
    try:
        vectorizer = TfidfVectorizer(stop_words='english', ngram_range=(1, 2))
        X = vectorizer.fit_transform(keywords)
        
        kmeans = KMeans(n_clusters=n_clusters, random_state=42, n_init=10)
        clusters = kmeans.fit_predict(X)
        
        clustered_keywords = {}
        for i, keyword in enumerate(keywords):
            cluster_id = clusters[i]
            if cluster_id not in clustered_keywords:
                clustered_keywords[cluster_id] = []
            clustered_keywords[cluster_id].append(keyword)
        
        return clustered_keywords
    except Exception as e:
        st.error(f"Clustering error: {str(e)}")
        return {0: keywords}  # Return all keywords in one cluster as fallback

def export_to_csv(data, filename):
    """Export data to CSV format"""
    df = pd.DataFrame(data)
    csv = df.to_csv(index=False)
    return csv.encode('utf-8')

def export_to_docx(content, filename):
    """Export content to DOCX format"""
    doc = docx.Document()
    
    for item in content:
        if isinstance(item, dict):
            if 'title' in item:
                doc.add_heading(item['title'], level=1)
            if 'content' in item:
                doc.add_paragraph(item['content'])
        else:
            doc.add_paragraph(str(item))
    
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

def export_to_pdf(content, filename):
    """Export content to PDF format"""
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter)
    styles = getSampleStyleSheet()
    story = []
    
    for item in content:
        if isinstance(item, dict):
            if 'title' in item:
                story.append(Paragraph(item['title'], styles['Heading1']))
                story.append(Spacer(1, 12))
            if 'content' in item:
                story.append(Paragraph(item['content'], styles['Normal']))
                story.append(Spacer(1, 12))
        else:
            story.append(Paragraph(str(item), styles['Normal']))
            story.append(Spacer(1, 12))
    
    doc.build(story)
    return buffer.getvalue()

# Main App
def main():
    # Initialize
    init_session_state()
    
    # Sidebar Configuration
    with st.sidebar:
        st.header("🔧 Configuration")
        
        # Quick Stats
        st.subheader("📊 Session Stats")
        col1, col2 = st.columns(2)
        with col1:
            st.metric("Content Generated", len(st.session_state.generated_content))
            st.metric("Keywords Found", len(st.session_state.discovered_keywords))
        with col2:
            st.metric("Clusters Created", len(st.session_state.keyword_clusters))
            st.metric("Metadata Sets", len(st.session_state.metadata_results))
    
    # Header
    st.markdown('<div class="main-header">🚀 AI SEO Content Automation Platform</div>', unsafe_allow_html=True)
    
    # Main Tabs
    tab1, tab2, tab3, tab4, tab5, tab6 = st.tabs([
        "📝 Content Generation", 
        "🎯 Metadata Optimizer", 
        "🔍 Keyword Discovery", 
        "📊 Keyword Clustering", 
        "📁 Export Manager",
        "📈 Dashboard"
    ])
    
    # Tab 1: Content Generation
    with tab1:
        st.markdown('<div class="module-header">📝 Create High-Converting SEO Content</div>', unsafe_allow_html=True)
        st.caption("Generate blog posts, product pages, and landing pages optimized for search engines and conversions.")
        
        # Content Type Selection
        content_type = st.selectbox(
            "Content Type",
            ["Blog Post", "Product Page", "Landing Page", "Pillar Page"]
        )
        
        col1, col2 = st.columns([2, 1])
        
        with col1:
            primary_keyword = st.text_input("Primary Keyword *", placeholder="e.g., AI marketing tools")
            secondary_keywords = st.text_input("Secondary Keywords", placeholder="marketing automation, AI tools, digital marketing")
            
            col1a, col1b = st.columns(2)
            with col1a:
                word_count = st.slider("Word Count", 300, 3000, 1500)
                tone = st.selectbox("Tone", ["Professional", "Casual", "Technical", "Conversational", "Authoritative"])
            with col1b:
                industry = st.text_input("Industry", placeholder="e.g., SaaS, E-commerce, Healthcare")
            
            additional_context = st.text_area("Additional Context (Optional)", placeholder="Any specific requirements, target audience details, or competitor information...")
        
        with col2:
            st.info("💡 **Tips for Better Content:**\n\n"
                   "• Use specific, long-tail primary keywords\n"
                   "• Include 3-5 related secondary keywords\n"
                   "• Provide industry context for better targeting\n"
                   "• Add competitor insights in additional context")
        
        # Validation and Generation
        if st.button("🚀 Generate Content", type="primary"):
            inputs = {
                'primary_keyword': primary_keyword,
                'secondary_keywords': secondary_keywords,
                'word_count': word_count,
                'tone': tone,
                'industry': industry,
                'additional_context': additional_context
            }
            
            # Apply fallbacks
            inputs = apply_fallbacks(inputs)
            
            # Show warnings for missing inputs
            warnings = validate_inputs({
                'primary_keyword': 'Primary keyword',
                'industry': 'Industry specification'
            }, inputs)
            
            if warnings:
                for warning in warnings:
                    st.warning(warning)
            
            # Generate content
            with st.spinner("Generating content..."):
                if content_type == "Blog Post":
                    prompt = PromptTemplates.blog_post_template(
                        inputs['primary_keyword'], inputs['secondary_keywords'],
                        inputs['word_count'], inputs['tone'], inputs['industry'],
                        inputs['additional_context']
                    )
                elif content_type == "Product Page":
                    prompt = PromptTemplates.product_page_template(
                        inputs['primary_keyword'], inputs['industry'],
                        inputs['secondary_keywords'], inputs['primary_keyword']
                    )
                else:
                    prompt = f"Create a {content_type.lower()} about {inputs['primary_keyword']} for the {inputs['industry']} industry. Tone: {inputs['tone']}. Word count: {inputs['word_count']}."
                
                # Use ask_gpt_insights for content generation
                result = ask_gpt_insights(summary="", user_question=prompt, chat_history="")
                
                if result and not result.startswith("Error") and not result.startswith("Please configure"):
                    # Store in session state
                    content_item = {
                        'type': content_type,
                        'title': f"{content_type}: {inputs['primary_keyword']}",
                        'content': result,
                        'timestamp': datetime.now().strftime("%Y-%m-%d %H:%M"),
                        'metadata': inputs
                    }
                    st.session_state.generated_content.append(content_item)
                    
                    st.success("✅ Content generated successfully!")
                    st.markdown("### Generated Content")
                    st.markdown(result)
                else:
                    st.error(result)
    
    # Tab 2: Metadata Optimizer
    with tab2:
        st.markdown('<div class="module-header">🎯 Smart Metadata Optimization</div>', unsafe_allow_html=True)
        st.caption("Generate SEO-perfect titles, descriptions, and headers that drive clicks and rankings.")
        
        # Input Method
        input_method = st.radio("Input Method", ["Paste Content", "Enter Topic/URL"])
        
        col1, col2 = st.columns([2, 1])
        
        with col1:
            if input_method == "Paste Content":
                content_input = st.text_area("Content to Optimize", height=200, placeholder="Paste your existing content here...")
            else:
                content_input = st.text_input("Topic or URL", placeholder="Enter topic or URL to optimize")
            
            col1a, col1b = st.columns(2)
            with col1a:
                meta_primary_keyword = st.text_input("Primary Keyword", placeholder="e.g., best CRM software")
                search_intent = st.selectbox("Search Intent", ["Informational", "Commercial", "Navigational", "Transactional"])
            with col1b:
                meta_industry = st.text_input("Industry", placeholder="e.g., SaaS, Healthcare")
                include_local = st.checkbox("Include Local SEO")
                if include_local:
                    location = st.text_input("Location", placeholder="e.g., New York, London")
        
        with col2:
            st.info("🎯 **Metadata Best Practices:**\n\n"
                   "• Title tags: 50-60 characters\n"
                   "• Meta descriptions: 150-160 characters\n"
                   "• Include primary keyword in title\n"
                   "• Use power words for better CTR")
        
        if st.button("🎯 Optimize Metadata", type="primary"):
            if content_input and meta_primary_keyword:
                inputs = {
                    'content': content_input,
                    'primary_keyword': meta_primary_keyword,
                    'search_intent': search_intent,
                    'industry': meta_industry or 'General'
                }
                
                with st.spinner("Optimizing metadata..."):
                    prompt = PromptTemplates.metadata_template(
                        inputs['content'], inputs['primary_keyword'],
                        inputs['search_intent'], inputs['industry']
                    )
                    
                    result = ask_gpt_insights(summary="", user_question=prompt, chat_history="")
                    
                    if result and not result.startswith("Error"):
                        metadata_item = {
                            'title': f"Metadata: {meta_primary_keyword}",
                            'content': result,
                            'timestamp': datetime.now().strftime("%Y-%m-%d %H:%M"),
                            'metadata': inputs
                        }
                        st.session_state.metadata_results.append(metadata_item)
                        
                        st.success("✅ Metadata optimized successfully!")
                        st.markdown("### Optimized Metadata")
                        st.markdown(result)
                    else:
                        st.error(result)
            else:
                st.error("Please provide content and primary keyword.")
    
    # Tab 3: Keyword Discovery
    with tab3:
        st.markdown('<div class="module-header">🔍 AI-Powered Keyword Research</div>', unsafe_allow_html=True)
        st.caption("Uncover profitable long-tail keywords your competitors are missing.")
        
        col1, col2 = st.columns([2, 1])
        
        with col1:
            seed_keyword = st.text_input("Seed Keyword *", placeholder="e.g., project management software")
            
            col1a, col1b = st.columns(2)
            with col1a:
                kw_industry = st.text_input("Industry", placeholder="e.g., SaaS, Marketing")
                target_audience = st.text_input("Target Audience", placeholder="e.g., small business owners, developers")
            with col1b:
                keyword_count = st.selectbox("Keyword Count", [25, 50, 100])
                include_local_kw = st.checkbox("Include Local Keywords")
                if include_local_kw:
                    kw_location = st.text_input("Location", placeholder="e.g., California, UK")
            
            # Intent Filters
            st.subheader("Intent Filters")
            col1a, col1b, col1c = st.columns(3)
            with col1a:
                info_intent = st.checkbox("Informational", value=True)
                commercial_intent = st.checkbox("Commercial", value=True)
            with col1b:
                question_based = st.checkbox("Question-based", value=True)
                comparison = st.checkbox("Comparison")
            with col1c:
                local_intent = st.checkbox("Local")
                navigational = st.checkbox("Navigational")
        
        with col2:
            st.info("🔍 **Research Tips:**\n\n"
                   "• Use specific seed keywords\n"
                   "• Define your target audience\n"
                   "• Mix different intent types\n"
                   "• Consider local variations")
        
        if st.button("🔍 Discover Keywords", type="primary"):
            if seed_keyword:
                inputs = {
                    'seed_keyword': seed_keyword,
                    'industry': kw_industry or 'General',
                    'audience': target_audience or 'General audience',
                    'include_local': include_local_kw,
                    'location': kw_location if include_local_kw else ''
                }
                
                with st.spinner("Discovering keywords..."):
                    prompt = PromptTemplates.keyword_discovery_template(
                        inputs['seed_keyword'], inputs['industry'],
                        inputs['audience'], inputs['include_local'], inputs['location']
                    )
                    
                    result = ask_gpt_insights(summary="", user_question=prompt, chat_history="")
                    
                    if result and not result.startswith("Error"):
                        # Parse keywords from result
                        lines = result.split('\n')
                        keywords = []
                        for line in lines:
                            if '|' in line and not line.startswith('Keyword'):
                                parts = line.split('|')
                                if len(parts) >= 4:
                                    keyword = parts[0].strip()
                                    if keyword and len(keyword.split()) > 1:
                                        keywords.append({
                                            'keyword': keyword,
                                            'intent': parts[1].strip(),
                                            'difficulty': parts[2].strip(),
                                            'content_type': parts[3].strip()
                                        })
                        
                        st.session_state.discovered_keywords = keywords
                        
                        st.success(f"✅ Discovered {len(keywords)} keywords!")
                        
                        # Display results
                        if keywords:
                            df = pd.DataFrame(keywords)
                            st.markdown("### Discovered Keywords")
                            st.dataframe(df, use_container_width=True)
                        else:
                            st.markdown("### Raw Results")
                            st.markdown(result)
                    else:
                        st.error(result)
            else:
                st.error("Please provide a seed keyword.")
    
    # Tab 4: Keyword Clustering
    with tab4:
        st.markdown('<div class="module-header">📊 Smart Keyword Clustering</div>', unsafe_allow_html=True)
        st.caption("Group related keywords into content clusters for maximum SEO impact.")
        
        col1, col2 = st.columns([2, 1])
        
        with col1:
            # Import options
            st.subheader("Import Keywords")
            import_option = st.radio("Import Method", ["From Keyword Discovery", "Upload CSV", "Paste Keywords"])
            
            keywords_to_cluster = []
            
            if import_option == "From Keyword Discovery":
                if st.session_state.discovered_keywords:
                    keywords_to_cluster = [kw['keyword'] for kw in st.session_state.discovered_keywords]
                    st.success(f"Loaded {len(keywords_to_cluster)} keywords from discovery")
                else:
                    st.warning("No keywords found. Please run keyword discovery first.")
            
            elif import_option == "Upload CSV":
                uploaded_file = st.file_uploader("Choose CSV file", type="csv")
                if uploaded_file:
                    df = pd.read_csv(uploaded_file)
                    if 'keyword' in df.columns:
                        keywords_to_cluster = df['keyword'].tolist()
                    else:
                        st.error("CSV must have a 'keyword' column")
            
            else:  # Paste Keywords
                keyword_text = st.text_area("Keywords (one per line)", height=200)
                if keyword_text:
                    keywords_to_cluster = [kw.strip() for kw in keyword_text.split('\n') if kw.strip()]
            
            # Clustering settings
            if keywords_to_cluster:
                st.subheader("Clustering Settings")
                col1a, col1b = st.columns(2)
                with col1a:
                    max_clusters = st.slider("Max Clusters", 3, min(20, len(keywords_to_cluster)), 5)
                with col1b:
                    similarity_threshold = st.selectbox("Similarity Threshold", ["Low", "Medium", "High"])
        
        with col2:
            st.info("📊 **Clustering Benefits:**\n\n"
                   "• Organize content strategy\n"
                   "• Identify content gaps\n"
                   "• Plan pillar pages\n"
                   "• Optimize internal linking")
        
        if keywords_to_cluster and st.button("📊 Create Clusters", type="primary"):
            with st.spinner("Creating keyword clusters..."):
                clusters = keyword_clustering(keywords_to_cluster, max_clusters)
                
                if clusters:
                    st.session_state.keyword_clusters = clusters
                    st.success(f"✅ Created {len(clusters)} keyword clusters!")
                    
                    # Display clusters
                    st.markdown("### Keyword Clusters")
                    for cluster_id, keywords in clusters.items():
                        with st.expander(f"Cluster {cluster_id + 1} ({len(keywords)} keywords)"):
                            st.write("**Primary keyword (suggested):**", keywords[0])
                            st.write("**Related keywords:**")
                            for kw in keywords[1:]:
                                st.write(f"• {kw}")
                else:
                    st.error("Failed to create clusters")
    
    # Tab 5: Export Manager
    with tab5:
        st.markdown('<div class="module-header">📁 Export Your SEO Assets</div>', unsafe_allow_html=True)
        st.caption("Download your optimized content in multiple formats, ready for publication.")
        
        # Content Selection
        st.subheader("Content Selection")
        
        export_items = []
        
        # Generated Content
        if st.session_state.generated_content:
            st.write("**Generated Content:**")
            for i, item in enumerate(st.session_state.generated_content):
                if st.checkbox(f"📝 {item['title']}", key=f"content_{i}"):
                    export_items.append(item)
        
        # Keywords
        if st.session_state.discovered_keywords:
            if st.checkbox("🔍 Discovered Keywords"):
                export_items.append({
                    'title': 'Discovered Keywords',
                    'content': pd.DataFrame(st.session_state.discovered_keywords).to_string(),
                    'type': 'keywords'
                })
        
        # Clusters
        if st.session_state.keyword_clusters:
            if st.checkbox("📊 Keyword Clusters"):
                clusters_content = ""
                for cluster_id, keywords in st.session_state.keyword_clusters.items():
                    clusters_content += f"\n\nCluster {cluster_id + 1}:\n"
                    clusters_content += "\n".join([f"• {kw}" for kw in keywords])
                
                export_items.append({
                    'title': 'Keyword Clusters',
                    'content': clusters_content,
                    'type': 'clusters'
                })
        
        # Metadata
        if st.session_state.metadata_results:
            st.write("**Metadata Results:**")
            for i, item in enumerate(st.session_state.metadata_results):
                if st.checkbox(f"🎯 {item['title']}", key=f"meta_{i}"):
                    export_items.append(item)
        
        # Export Options
        if export_items:
            st.subheader("Export Options")
            col1, col2 = st.columns(2)
            
            with col1:
                export_format = st.selectbox("Format", ["CSV", "DOCX", "PDF", "JSON"])
                include_metadata = st.checkbox("Include Metadata", value=True)
            
            with col2:
                include_timestamp = st.checkbox("Include Timestamps", value=True)
                include_seo_checklist = st.checkbox("Include SEO Checklist")
            
            # Export buttons
            col1, col2, col3 = st.columns(3)
            
            with col1:
                if st.button("📁 Export Selected", type="primary"):
                    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                    filename = f"seo_export_{timestamp}"
                    
                    try:
                        if export_format == "CSV":
                            # Convert to CSV format
                            export_data = []
                            for item in export_items:
                                export_data.append({
                                    'Title': item['title'],
                                    'Type': item.get('type', 'content'),
                                    'Content': item['content'][:1000] + "..." if len(item['content']) > 1000 else item['content'],
                                    'Timestamp': item.get('timestamp', '')
                                })
                            
                            csv_data = export_to_csv(export_data, filename)
                            st.download_button(
                                label="📥 Download CSV",
                                data=csv_data,
                                file_name=f"{filename}.csv",
                                mime="text/csv"
                            )
                        
                        elif export_format == "DOCX":
                            docx_data = export_to_docx(export_items, filename)
                            st.download_button(
                                label="📥 Download DOCX",
                                data=docx_data,
                                file_name=f"{filename}.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                            )
                        
                        elif export_format == "PDF":
                            pdf_data = export_to_pdf(export_items, filename)
                            st.download_button(
                                label="📥 Download PDF",
                                data=pdf_data,
                                file_name=f"{filename}.pdf",
                                mime="application/pdf"
                            )
                        
                        elif export_format == "JSON":
                            json_data = json.dumps(export_items, indent=2, default=str)
                            st.download_button(
                                label="📥 Download JSON",
                                data=json_data.encode('utf-8'),
                                file_name=f"{filename}.json",
                                mime="application/json"
                            )
                        
                        st.success("✅ Export prepared successfully!")
                        
                    except Exception as e:
                        st.error(f"Export failed: {str(e)}")
            
            with col2:
                if st.button("📋 Preview"):
                    st.markdown("### Export Preview")
                    for item in export_items:
                        st.markdown(f"**{item['title']}**")
                        st.markdown(item['content'][:500] + "..." if len(item['content']) > 500 else item['content'])
                        st.markdown("---")
            
            with col3:
                if st.button("🗑️ Clear Selection"):
                    st.rerun()
        
        else:
            st.info("No content selected for export. Generate some content first!")
    
    # Tab 6: Dashboard
    with tab6:
        st.markdown('<div class="module-header">📈 Dashboard & Analytics</div>', unsafe_allow_html=True)
        
        # Quick Stats
        col1, col2, col3, col4 = st.columns(4)
        
        with col1:
            st.markdown('<div class="metric-card">', unsafe_allow_html=True)
            st.metric("Content Generated", len(st.session_state.generated_content))
            st.markdown('</div>', unsafe_allow_html=True)
        
        with col2:
            st.markdown('<div class="metric-card">', unsafe_allow_html=True)
            st.metric("Keywords Found", len(st.session_state.discovered_keywords))
            st.markdown('</div>', unsafe_allow_html=True)
        
        with col3:
            st.markdown('<div class="metric-card">', unsafe_allow_html=True)
            st.metric("Clusters Created", len(st.session_state.keyword_clusters))
            st.markdown('</div>', unsafe_allow_html=True)
        
        with col4:
            st.markdown('<div class="metric-card">', unsafe_allow_html=True)
            st.metric("Metadata Sets", len(st.session_state.metadata_results))
            st.markdown('</div>', unsafe_allow_html=True)
        
        # Recent Activity
        st.subheader("📋 Recent Activity")
        
        all_activities = []
        
        # Add content generation activities
        for item in st.session_state.generated_content[-5:]:
            all_activities.append({
                'timestamp': item.get('timestamp', ''),
                'activity': f"📝 Generated {item['type']}: {item['title']}",
                'type': 'content'
            })
        
        # Add keyword discovery activities
        if st.session_state.discovered_keywords:
            all_activities.append({
                'timestamp': datetime.now().strftime("%Y-%m-%d %H:%M"),
                'activity': f"🔍 Discovered {len(st.session_state.discovered_keywords)} keywords",
                'type': 'keywords'
            })
        
        # Add clustering activities
        if st.session_state.keyword_clusters:
            all_activities.append({
                'timestamp': datetime.now().strftime("%Y-%m-%d %H:%M"),
                'activity': f"📊 Created {len(st.session_state.keyword_clusters)} keyword clusters",
                'type': 'clusters'
            })
        
        # Add metadata activities
        for item in st.session_state.metadata_results[-3:]:
            all_activities.append({
                'timestamp': item.get('timestamp', ''),
                'activity': f"🎯 Optimized metadata: {item['title']}",
                'type': 'metadata'
            })
        
        # Sort by timestamp and display
        all_activities.sort(key=lambda x: x['timestamp'], reverse=True)
        
        if all_activities:
            for activity in all_activities[:10]:  # Show last 10 activities
                st.write(f"• {activity['activity']}")
                if activity['timestamp']:
                    st.caption(f"  {activity['timestamp']}")
        else:
            st.info("No recent activity. Start by generating some content!")
        
        # Quick Actions
        st.subheader("⚡ Quick Actions")
        col1, col2, col3, col4 = st.columns(4)
        
        with col1:
            if st.button("🚀 New Content", use_container_width=True):
                st.switch_page("tab1")  # This would need to be implemented with session state
        
        with col2:
            if st.button("🔍 Keyword Research", use_container_width=True):
                st.switch_page("tab3")
        
        with col3:
            if st.button("📊 View Clusters", use_container_width=True):
                if st.session_state.keyword_clusters:
                    st.markdown("### Current Clusters")
                    for cluster_id, keywords in st.session_state.keyword_clusters.items():
                        st.write(f"**Cluster {cluster_id + 1}:** {', '.join(keywords[:3])}...")
                else:
                    st.info("No clusters created yet.")
        
        with col4:
            if st.button("📁 Recent Exports", use_container_width=True):
                st.info("Export history feature coming soon!")
        
        # Content Analysis
        if st.session_state.generated_content or st.session_state.discovered_keywords:
            st.subheader("📊 Content Analysis")
            
            # Word count analysis
            if st.session_state.generated_content:
                word_counts = []
                content_types = []
                
                for item in st.session_state.generated_content:
                    word_count = len(item['content'].split())
                    word_counts.append(word_count)
                    content_types.append(item.get('type', 'Unknown'))
                
                if word_counts:
                    df_analysis = pd.DataFrame({
                        'Content Type': content_types,
                        'Word Count': word_counts
                    })
                    
                    col1, col2 = st.columns(2)
                    
                    with col1:
                        st.markdown("**Content by Type**")
                        type_counts = pd.Series(content_types).value_counts().reset_index()
                        type_counts.columns = ['Type', 'Count']
                        st.bar_chart(type_counts.set_index('Type'))
                    
                    with col2:
                        st.markdown("**Word Count Distribution**")
                        st.line_chart(pd.DataFrame({'Word Count': word_counts}))
            
            # Keyword analysis
            if st.session_state.discovered_keywords:
                st.markdown("**Keyword Intent Distribution**")
                intents = [kw.get('intent', 'Unknown') for kw in st.session_state.discovered_keywords]
                intent_counts = pd.Series(intents).value_counts().reset_index()
                intent_counts.columns = ['Intent', 'Count']
                st.bar_chart(intent_counts.set_index('Intent'))
        
        # System Information
        with st.expander("🔧 System Information"):
            st.write(f"**Current AI Provider:** {st.session_state.api_provider}")
            st.write(f"**Session Started:** {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            st.write(f"**Total Items in Memory:** {len(st.session_state.generated_content) + len(st.session_state.discovered_keywords) + len(st.session_state.keyword_clusters) + len(st.session_state.metadata_results)}")
            
            if st.button("🗑️ Clear All Data"):
                st.session_state.generated_content = []
                st.session_state.discovered_keywords = []
                st.session_state.keyword_clusters = []
                st.session_state.metadata_results = []
                st.success("All data cleared!")
                st.rerun()

# Additional utility functions for enhanced functionality
def calculate_keyword_density(content, keyword):
    """Calculate keyword density in content"""
    content_lower = content.lower()
    keyword_lower = keyword.lower()
    word_count = len(content.split())
    keyword_count = content_lower.count(keyword_lower)
    return (keyword_count / word_count) * 100 if word_count > 0 else 0

def suggest_improvements(content, primary_keyword=""):
    """Suggest content improvements"""
    suggestions = []
    word_count = len(content.split())
    
    # Word count suggestions
    if word_count < 300:
        suggestions.append("⚠️ Content is quite short. Consider expanding to at least 300 words for better SEO.")
    elif word_count > 3000:
        suggestions.append("📝 Very long content. Consider breaking into multiple pieces or adding subheadings.")
    
    # Keyword density
    if primary_keyword:
        density = calculate_keyword_density(content, primary_keyword)
        if density < 0.5:
            suggestions.append(f"🔍 Low keyword density ({density:.1f}%). Consider adding '{primary_keyword}' 2-3 more times.")
        elif density > 3:
            suggestions.append(f"⚠️ High keyword density ({density:.1f}%). Reduce keyword usage to avoid stuffing.")
    
    # Structure suggestions
    if content.count('\n## ') < 2:
        suggestions.append("📋 Add more H2 headings to improve content structure and readability.")
    
    return suggestions

def generate_seo_checklist():
    """Generate SEO checklist for content"""
    return """
    ## SEO Content Checklist ✅
    
    ### Technical SEO
    - [ ] Title tag includes primary keyword (50-60 chars)
    - [ ] Meta description is compelling (150-160 chars)
    - [ ] H1 tag contains primary keyword
    - [ ] URL is SEO-friendly and readable
    - [ ] Images have descriptive alt text
    
    ### Content Quality
    - [ ] Content is 300+ words
    - [ ] Primary keyword density is 0.5-2.5%
    - [ ] Secondary keywords used naturally
    - [ ] Content provides value to readers
    - [ ] Includes relevant internal/external links
    
    ### User Experience
    - [ ] Content is scannable with subheadings
    - [ ] Paragraphs are short (3-4 sentences max)
    - [ ] Includes bullet points or numbered lists
    - [ ] Has a clear call-to-action
    - [ ] Mobile-friendly formatting
    
    ### Advanced SEO
    - [ ] Schema markup implemented
    - [ ] Featured snippet optimization
    - [ ] Related keywords included
    - [ ] Content freshness maintained
    - [ ] Social sharing optimized
    """

# Error handling and logging
def log_error(error_message, context=""):
    """Log errors for debugging"""
    timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    error_log = f"[{timestamp}] {context}: {error_message}"
    
    # In a production environment, you'd want to log this to a file or service
    print(error_log)  # For now, just print to console

def ask_gpt_insights(summary, user_question, chat_history=""):
    load_dotenv()
    endpoint = os.getenv("AZURE_OPENAI_ENDPOINT")
    deployment = os.getenv("AZURE_OPENAI_DEPLOYMENT")
    api_key = os.getenv("AZURE_OPENAI_API_KEY")
    api_version = os.getenv("AZURE_OPENAI_API_VERSION")

    client = AzureOpenAI(
        azure_endpoint=endpoint,
        api_key=api_key,
        api_version=api_version,
    )

    # Compose the chat history/messages
    messages = []
    if chat_history:
        messages.append({"role": "system", "content": chat_history})
    if summary:
        messages.append({"role": "system", "content": f"Summary: {summary}"})
    messages.append({"role": "user", "content": user_question})

    # Call the Azure OpenAI chat completion endpoint
    response = client.chat.completions.create(
        model=deployment,
        messages=messages,
        max_tokens=1024,
        temperature=0.7
    )
    return response.choices[0].message.content

# Run the main application
if __name__ == "__main__":
    try:
        main()
    except Exception as e:
        st.error("An unexpected error occurred. Please refresh the page and try again.")
        log_error(str(e), "Main application")
        st.exception(e)  # Remove this in production